from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.infrastructure.parsers.document_parser import PyMuPDFParser, get_parser


@pytest.fixture
def parser():
    return PyMuPDFParser()


def test_parse_text_file(tmp_path: Path, parser: PyMuPDFParser):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("This is a sample clinical protocol.", encoding="utf-8")

    pages = parser.parse(file_path, "text/plain")

    assert len(pages) == 1
    assert pages[0]["text"] == "This is a sample clinical protocol."
    assert pages[0]["page_number"] == 1


def test_parse_docx_file(tmp_path: Path, parser: PyMuPDFParser):
    file_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(file_path))

    pages = parser.parse(
        file_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    assert len(pages) == 1
    assert "First paragraph." in pages[0]["text"]
    assert "Second paragraph." in pages[0]["text"]


def test_parse_unsupported_content_type(tmp_path: Path, parser: PyMuPDFParser):
    file_path = tmp_path / "sample.xyz"
    file_path.write_text("content")

    with pytest.raises(ValueError, match="Unsupported content type"):
        parser.parse(file_path, "application/unknown")


def test_get_parser_returns_parser():
    parser = get_parser()
    assert isinstance(parser, PyMuPDFParser)
